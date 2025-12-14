"""
Word Document (DOCX) Extractor

Extracts text, tables, and metadata from Word documents.
"""

from pathlib import Path
from typing import Dict, List, Any, Optional, TYPE_CHECKING
from io import BytesIO
import re

if TYPE_CHECKING:
    from docx.api import Document as DocxDocumentType
else:
    DocxDocumentType = Any  # type: ignore

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore
from PIL import Image
import pytesseract

from ..utils import (
    detect_performance_blocks,
    detect_source_notes,
    detect_language_safe,
    detect_disclaimer_flags,
    detect_glossary,
    detect_legal_notice,
    categorize_disclaimer,
    extract_title_fields,
    extract_identifiers,
    extract_issuer_mentions,
    analyze_performance_sentence,
    parse_table_source,
    detect_countries,
    country_entries,
    normalize_table_rows,
    extract_docx_table,
    prepare_image_for_ocr,
    extract_text_from_image,
)
from ...parsers.registration import COUNTRY_PATTERNS


def extract_docx(
    file_path: Path,
    tesseract_lang: str = "eng",
    tesseract_config: str = "--psm 6"
) -> Dict[str, Any]:
    """
    Extract text and tables from Word document.
    
    Args:
        file_path: Path to DOCX file
        tesseract_lang: Tesseract language code
        tesseract_config: Tesseract configuration
        
    Returns:
        Dict with extracted data
    """
    if DocxDocument is None:
        raise ImportError("python-docx not installed")
    
    doc = DocxDocument(str(file_path))
    
    full_text = []
    paragraphs_data = []
    tables_data = []
    paragraph_summaries = []
    performance_sections = []
    disclaimer_paragraphs = []
    glossary_paragraphs = []
    country_mentions = set()
    table_sources = []
    identifiers: List[Dict[str, Any]] = []
    country_entries_list: List[Dict[str, Any]] = []
    issuer_mentions: List[Dict[str, Any]] = []
    legal_notice_paragraph = None
    performance_table_entries: List[Dict[str, Any]] = []

    # Extract paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            full_text.append(text)
            paragraphs_data.append({
                'paragraph_number': para_idx,
                'text': text,
                'style': paragraph.style.name if paragraph.style else None
            })

            # Compile regex patterns for detect_countries
            compiled_patterns = {country: re.compile(pattern, re.IGNORECASE) for country, pattern in COUNTRY_PATTERNS.items()}
            paragraph_countries = detect_countries(text, compiled_patterns)
            country_mentions.update(paragraph_countries)

            heading = paragraph.style.name if paragraph.style else None
            country_entries_list.extend(
                country_entries(
                    paragraph_countries,
                    heading,
                    text,
                    {'paragraph_number': para_idx}
                )
            )

            issuer_mentions.extend(
                extract_issuer_mentions(
                    text,
                    {'paragraph_number': para_idx, 'heading': heading}
                )
            )

            perf_sentences = detect_performance_blocks(text)
            if perf_sentences:
                entries = [analyze_performance_sentence(sentence) for sentence in perf_sentences]
                performance_sections.append({
                    'paragraph_number': para_idx,
                    'sentences': perf_sentences,
                    'entries': entries
                })

            if detect_disclaimer_flags(text):
                disclaimer_paragraphs.append(para_idx)

            if detect_glossary(text):
                glossary_paragraphs.append(para_idx)

            table_sources.extend({
                'paragraph_number': para_idx,
                **parse_table_source(note)
            } for note in detect_source_notes(text))

            if para_idx == 1:
                title_info = extract_title_fields(text)

            identifier_entries = extract_identifiers(
                text,
                {'paragraph_number': para_idx}
            )
            if identifier_entries:
                identifiers.extend(identifier_entries)

            if legal_notice_paragraph is None and detect_legal_notice(text):
                legal_notice_paragraph = para_idx

            paragraph_summaries.append({
                'paragraph_number': para_idx,
                'style': paragraph.style.name if paragraph.style else None,
                'contains_disclaimer': para_idx in disclaimer_paragraphs,
                'disclaimer_categories': categorize_disclaimer(text) if para_idx in disclaimer_paragraphs else [],
                'contains_performance': bool(perf_sentences),
                'contains_glossary': para_idx in glossary_paragraphs,
                'countries': paragraph_countries,
                'language': detect_language_safe(text),
            })

    # Extract tables
    for table_idx, table in enumerate(doc.tables, 1):
        table_data = extract_docx_table(table)
        if table_data:
            tables_data.append({
                'table_index': table_idx,
                'data': table_data
            })
            for entry in normalize_table_rows(table_data):
                entry.update({'table_index': table_idx})
                performance_table_entries.append(entry)
     
    # Extract images and perform OCR
    image_texts = _extract_docx_images(doc, tesseract_lang, tesseract_config)
    if image_texts:
        full_text.extend(image_texts)
     
    return {
        'text': '\n'.join(full_text),
        'paragraphs': paragraphs_data,
        'tables': tables_data,
        'image_ocr_text': image_texts if image_texts else None,
        'paragraph_summaries': paragraph_summaries,
        'performance_sections': performance_sections,
        'disclaimer_paragraphs': disclaimer_paragraphs,
        'glossary_paragraphs': glossary_paragraphs,
        'table_sources': table_sources,
        'structure': {
            'disclaimer_paragraphs': disclaimer_paragraphs,
            'performance_paragraphs': [section['paragraph_number'] for section in performance_sections],
            'glossary_paragraphs': glossary_paragraphs,
            'countries_detected': sorted(country_mentions),
            'disclaimer_categories': {
                para['paragraph_number']: para['disclaimer_categories']
                for para in paragraph_summaries if para['disclaimer_categories']
            },
            'country_entries': country_entries_list,
            'issuer_mentions': issuer_mentions,
            'has_glossary': bool(glossary_paragraphs),
            'has_management_notice': legal_notice_paragraph is not None,
            'legal_notice_paragraph': legal_notice_paragraph,
        },
        'identifiers': identifiers,
        'country_entries': country_entries_list,
        'issuer_mentions': issuer_mentions,
        'performance_table_entries': performance_table_entries,
        'title_information': title_info if 'title_info' in locals() else None,
        'total_paragraphs': len(paragraphs_data),
        'total_tables': len(tables_data),
        'ocr': bool(image_texts)
    }


def _extract_docx_images(
    doc: Any,  # type: ignore[valid-type]
    tesseract_lang: str = "eng",
    tesseract_config: str = "--psm 6"
) -> List[str]:
    """Perform OCR on inline images within a DOCX document."""
    if pytesseract is None or Image is None:
        return []

    ocr_texts = []
    for rel in doc.part.related_parts.values():
        content_type = getattr(rel, "content_type", "")
        if not content_type.startswith("image/"):
            continue
        try:
            with Image.open(BytesIO(rel.blob)) as img:
                img = prepare_image_for_ocr(img)
                text = extract_text_from_image(img, tesseract_lang, tesseract_config)
                if text:
                    ocr_texts.append(text)
        except Exception:
            continue
    return ocr_texts

