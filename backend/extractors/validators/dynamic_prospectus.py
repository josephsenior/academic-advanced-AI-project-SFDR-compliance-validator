"""
Dynamic Prospectus Extractor

Automatically detects and extracts subscription/redemption fee information
from DOCX prospectus files. Compares extracted fees with performance table
mentions to identify potential disclosure gaps.

Only extracts fees if they were mentioned as requirements (e.g., German Rule 3
or performance fee display rules).
"""

import logging
import re
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
from backend.extractors.rules.models import ComplianceIssue
from backend.extractors.rules.enums import ComplianceIssueType
from .base import BaseValidator

logger = logging.getLogger(__name__)


@dataclass
class FeeExtraction:
    """Extracted fee information from prospectus"""
    fee_type: str  # 'subscription', 'redemption', 'management', 'performance'
    fee_value: Optional[str]  # e.g., "1.5%", "max 1%", "up to 2%"
    fee_description: str
    location: str  # Where in prospectus found
    context: str  # Surrounding text for verification


class DynamicProspectusExtractor(BaseValidator):
    """
    Extracts subscription/redemption fees from prospectus documents.
    
    Designed to support compliance rules that require fee disclosures,
    such as German Rule 3 (fee loading for performance tables).
    """
    
    # Keywords indicating different fee types
    FEE_PATTERNS = {
        'subscription': {
            'FR': ['frais de souscription', 'frais d\'entrée', 'commission', 'frais initiaux'],
            'EN': ['subscription', 'entry fee', 'initial fee', 'upfront fee'],
            'DE': ['Souscriptions', 'Gebühren', 'Einmalgebühr', 'Eingangsgebühr']
        },
        'redemption': {
            'FR': ['frais de rétraction', 'frais de sortie', 'frais de remboursement', 'frais de rachat'],
            'EN': ['redemption', 'exit fee', 'withdrawal fee', 'closing fee'],
            'DE': ['Rücknahme', 'Ausstiegsgebühren', 'Rücknahmegebühr', 'Auszahlungsgebühr']
        },
        'management': {
            'FR': ['frais de gestion', 'frais annuels', 'ocf'],
            'EN': ['management fee', 'annual fee', 'ocf', 'ongoing charges'],
            'DE': ['Verwaltungsgebühren', 'Jahresgebühren', 'laufende Gebühren']
        },
        'performance': {
            'FR': ['frais de performance', 'commission de surperformance', 'frais de performance'],
            'EN': ['performance fee', 'outperformance fee', 'performance charge'],
            'DE': ['Leistungsgebühr', 'Überrenditegebühr', 'Leistungsgebühr']
        }
    }
    
    def __init__(self, extract_management: bool = False, extract_performance: bool = False):
        """
        Initialize extractor.
        
        Args:
            extract_management: Also extract management fees (default: subscription/redemption only)
            extract_performance: Also extract performance fees
        """
        self.extract_management = extract_management
        self.extract_performance = extract_performance
    
    def validate(
        self,
        extraction_result: Dict[str, Any],
        metadata: Optional[Dict[str, Any]] = None,
        client_type: Optional[Any] = None,
        fund_type: Optional[Any] = None
    ) -> List[ComplianceIssue]:
        """
        Extract fees from prospectus and validate disclosure.
        
        Returns ComplianceIssue if fees are found in prospectus but
        not mentioned in performance tables (if they should be).
        """
        issues: List[ComplianceIssue] = []
        
        # Get prospectus path
        prospectus_path = self._get_prospectus_path(extraction_result, metadata)
        
        if not prospectus_path:
            logger.debug("No prospectus file found for fee extraction")
            return issues
        
        # Extract fees from prospectus
        prospectus_fees = self._extract_from_prospectus(prospectus_path, metadata)
        
        if not prospectus_fees:
            return issues
        
        # Check if fees are mentioned in performance tables
        performance_text = self._get_performance_text(extraction_result)
        
        # Validate that fees mentioned in prospectus are also in performance sections
        for fee in prospectus_fees:
            if fee.fee_type in ['subscription', 'redemption']:
                # These should be mentioned in performance tables
                if not self._fee_mentioned_in_text(fee, performance_text):
                    issue = ComplianceIssue(
                        issue_type=ComplianceIssueType.FEE_DISCLOSURE_INCOMPLETE,
                        issue_category='compliance',
                        severity='warning',
                        rule_reference="Fee Disclosure (Dynamic Prospectus)",
                        location=f"Performance Table (missing reference to {fee.fee_type} fee)",
                        message=(
                            f"{fee.fee_type.capitalize()} fee found in prospectus ({fee.fee_value}) "
                            f"but not mentioned in performance table"
                        ),
                        context=(
                            f"Prospectus location: {fee.location}\n"
                            f"Fee description: {fee.fee_description}\n"
                            f"Fee value: {fee.fee_value}"
                        ),
                        suggestion=(
                            f"Add reference to {fee.fee_type} fee ({fee.fee_value}) in first/last year "
                            f"of performance table or explicitly state if not applicable"
                        ),
                        auto_fixable=False
                    )
                    issues.append(issue)
        
        return issues
    
    def _get_prospectus_path(self, extraction_result: Dict[str, Any], metadata: Optional[Dict[str, Any]]) -> Optional[str]:
        """Find prospectus DOCX file."""
        # Check metadata
        if metadata and 'prospectus_path' in metadata:
            path = metadata['prospectus_path']
            if isinstance(path, str) and path.endswith('.docx'):
                return path
        
        # Look for DOCX files in same directory as PPTX
        if 'file_path' in extraction_result:
            pptx_path = Path(extraction_result['file_path'])
            docx_files = list(pptx_path.parent.glob('*.docx'))
            
            # Prefer "prospectus" or "document" in name
            for docx in docx_files:
                name = docx.name.lower()
                if 'prospectus' in name or 'règlement' in name or 'docum' in name:
                    return str(docx)
            
            # Otherwise take first DOCX
            if docx_files:
                return str(docx_files[0])
        
        return None
    
    def _extract_from_prospectus(self, docx_path: str, metadata: Optional[Dict[str, Any]]) -> List[FeeExtraction]:
        """
        Extract fee information from DOCX prospectus.
        
        Returns list of FeeExtraction objects.
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not available for prospectus extraction")
            return []
        
        extractions = []
        
        try:
            doc = Document(docx_path)
            
            # Determine document language
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            doc_lang = self._detect_language(doc_text, metadata)
            
            # Extract fees
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text
                
                # Look for each fee type
                for fee_type, lang_keywords in self.FEE_PATTERNS.items():
                    if fee_type == 'management' and not self.extract_management:
                        continue
                    if fee_type == 'performance' and not self.extract_performance:
                        continue
                    
                    keywords = lang_keywords.get(doc_lang, lang_keywords.get('EN', []))
                    
                    for keyword in keywords:
                        if keyword.lower() in text.lower():
                            # Extract fee value and details
                            fee_value = self._extract_fee_value(text)
                            context = self._get_surrounding_text(doc.paragraphs, para_idx, 2)
                            
                            extraction = FeeExtraction(
                                fee_type=fee_type,
                                fee_value=fee_value,
                                fee_description=text[:200],
                                location=f"Paragraph {para_idx + 1}",
                                context=context
                            )
                            
                            extractions.append(extraction)
                            break  # Only one extraction per paragraph
        
        except Exception as e:
            logger.warning(f"Error extracting from prospectus: {e}")
        
        return extractions
    
    def _detect_language(self, text: str, metadata: Optional[Dict[str, Any]]) -> str:
        """Detect document language."""
        if metadata and 'language_code' in metadata:
            lang = metadata['language_code'].upper()
            if lang in ['FR', 'EN', 'DE']:
                return lang
        
        # Simple detection
        text_lower = text.lower()
        fr_count = sum(1 for word in ['est', 'pour', 'dans'] if f' {word} ' in f' {text_lower} ')
        de_count = sum(1 for word in ['ist', 'für', 'mit'] if f' {word} ' in f' {text_lower} ')
        
        if fr_count > de_count:
            return 'FR'
        elif de_count > fr_count:
            return 'DE'
        
        return 'EN'
    
    def _extract_fee_value(self, text: str) -> Optional[str]:
        """
        Extract fee percentage or amount from text.
        
        Looks for patterns like "1.5%", "up to 2%", "max 1.5%", etc.
        """
        # Look for percentage
        percent_match = re.search(r'(?:up\s+to|max|maximum|jusqu\'?à|maximal)?\s*(\d+(?:[.,]\d+)?)\s*%', text, re.IGNORECASE)
        if percent_match:
            return f"{percent_match.group(1)}%"
        
        # Look for just percentage
        percent_match = re.search(r'(\d+(?:[.,]\d+)?)\s*%', text)
        if percent_match:
            return f"{percent_match.group(1)}%"
        
        # Look for flat amount
        amount_match = re.search(r'(?:EUR?|€)\s*(\d+(?:[.,]\d+)?)', text, re.IGNORECASE)
        if amount_match:
            return f"€{amount_match.group(1)}"
        
        return None
    
    def _get_surrounding_text(self, paragraphs, idx: int, context_paras: int = 2) -> str:
        """Get surrounding paragraphs for context."""
        start = max(0, idx - context_paras)
        end = min(len(paragraphs), idx + context_paras + 1)
        
        context_parts = []
        for i in range(start, end):
            if paragraphs[i].text.strip():
                context_parts.append(paragraphs[i].text[:100])
        
        return ' '.join(context_parts)
    
    def _get_performance_text(self, extraction_result: Dict[str, Any]) -> str:
        """Extract performance-related text from extraction result."""
        text_parts = []
        
        # Look for performance sections
        if 'performance_sections' in extraction_result:
            for perf_section in extraction_result['performance_sections']:
                if isinstance(perf_section, dict):
                    text_parts.append(perf_section.get('content', '') or perf_section.get('text', ''))
                else:
                    text_parts.append(str(perf_section))
        
        # Look for performance in slides
        if 'slides' in extraction_result:
            for slide in extraction_result['slides']:
                if isinstance(slide, dict):
                    text = slide.get('content', '') or slide.get('text', '')
                    if 'performance' in text.lower() or 'historical' in text.lower():
                        text_parts.append(text)
        
        # Main text
        if 'text' in extraction_result:
            text = extraction_result['text']
            if 'performance' in text.lower():
                text_parts.append(text)
        
        return ' '.join(filter(None, text_parts))
    
    def _fee_mentioned_in_text(self, fee: FeeExtraction, text: str) -> bool:
        """
        Check if a fee is mentioned in given text.
        
        Uses fee type and value to search for mention.
        """
        text_lower = text.lower()
        fee_type_lower = fee.fee_type.lower()
        
        # Check if fee type is mentioned
        if fee_type_lower not in text_lower:
            return False
        
        # Check if fee value is mentioned (if extracted)
        if fee.fee_value:
            # Normalize fee value for matching
            value = fee.fee_value.replace('%', '').strip()
            if value in text:
                return True
        
        return True  # Fee type mentioned is sufficient
