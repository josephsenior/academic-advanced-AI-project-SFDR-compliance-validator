from pathlib import Path
from docx import Document
import json
import re

DOCX_PATH = Path("dataset/Synthèse règles présentations commerciales.docx")
OUT_PATH = Path("data/slide_content_map.json")

def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip().lower()

def main():
    if not DOCX_PATH.exists():
        print("Docx file not found:", DOCX_PATH)
        return
    doc = Document(str(DOCX_PATH))
    # Heuristic: look for lines that start with slide number or "Page" or "Slide"
    slide_map = {}
    current_slide = None
    buffer = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        m = re.match(r'^(slide|page)\s*[:#]?\s*(\d+)\b', txt, re.I)
        m2 = re.match(r'^(\d+)\s*[-–]\s*(.+)$', txt)
        if m:
            # flush previous
            if current_slide and buffer:
                slide_map[current_slide] = normalize(" ".join(buffer))
            current_slide = int(m.group(2))
            buffer = [txt]
        elif m2:
            if current_slide and buffer:
                slide_map[current_slide] = normalize(" ".join(buffer))
            current_slide = int(m2.group(1))
            buffer = [m2.group(2)]
        else:
            # If the paragraph looks like a short description, append
            buffer.append(txt)

    if current_slide and buffer:
        slide_map[current_slide] = normalize(" ".join(buffer))

    # If nothing found, fallback: chunk document into sequential pages of 1..N
    if not slide_map:
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # split into groups of ~8 paragraphs per slide as heuristic
        chunk_size = 8
        for i in range(0, len(texts), chunk_size):
            slide_idx = (i // chunk_size) + 1
            slide_map[slide_idx] = normalize(" ".join(texts[i:i+chunk_size]))

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUT_PATH.write_text(json.dumps(slide_map, indent=2, ensure_ascii=False), encoding="utf-8")
    print("Wrote slide mapping to", OUT_PATH)

if __name__ == "__main__":
    main()

